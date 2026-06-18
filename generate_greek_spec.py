from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create Document
doc = Document()

# Set up styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title_run = title.add_run('TaskBoard Lite')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Τεχνική & Προδιαγραφή Έργου')
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(68, 114, 196)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Assignment info
assignment = doc.add_paragraph('Μάθημα: Διαχείριση Έργων Πληροφορικής')
assignment_run = assignment.runs[0]
assignment_run.font.size = Pt(11)
assignment_run.italic = True
assignment.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Date
date_para = doc.add_paragraph(f'Ημερομηνία Εγγράφου: {datetime.now().strftime("%d %B %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# TABLE OF CONTENTS
doc.add_heading('Πίνακας Περιεχομένων', level=1)
toc_items = [
    '1. Επισκόπηση Εργασίας',
    '2. Στόχοι Μάθησης',
    '3. Scope & Περιορισμοί Έργου',
    '4. Παραδοτέα',
    '5. Roadmap 3 Εβδομάδων',
    '6. Εκτελεστική Περίληψη',
    '7. Τεχνική Αρχιτεκτονική',
    '8. Μοντέλο Δεδομένων',
    '9. Τελικά Σημεία API',
    '10. Κατάσταση Υλοποίησης',
    '11. Κριτήρια Αξιολόγησης',
    '12. Εναλλακτικά Domains'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# SECTION 1: Assignment Overview
doc.add_heading('1. Επισκόπηση Εργασίας', level=1)
doc.add_paragraph(
    'Το TaskBoard Lite είναι μια διαδικτυακή εφαρμογή προσωπικής διαχείρισης εργασιών που '
    'υλοποιεί οργάνωση εργασιών σε στιλ Kanban (To Do / In Progress / Done), εμπνευσμένη από '
    'τις αρχές του Scrum για αποτελεσματική διαχείριση εργασιών. Αυτό είναι ένα ατομικό ακαδημαϊκό '
    'έργο για το μάθημα "Διαχείριση Έργων Πληροφορικής" με διάρκεια 3 εβδομάδων (~30-40 ώρες συνολικά).'
)

doc.add_heading('1.1 Τύπος Έργου', level=2)
doc.add_paragraph('Ατομική Εργασία')

doc.add_heading('1.2 Διάρκεια', level=2)
doc.add_paragraph('3 εβδομάδες (~30-40 συνολικές ώρες)')

doc.add_heading('1.3 Μορφή', level=2)
doc.add_paragraph('Πλήρης διαδικτυακή εφαρμογή με τεχνική τεκμηρίωση')

# SECTION 2: Learning Objectives
doc.add_heading('2. Στόχοι Μάθησης', level=1)
learning_objectives = [
    'Σχεδιασμός εφαρμογών χρησιμοποιώντας αρχιτεκτονική πολλαπλών επιπέδων (presentation / business / data)',
    'Υλοποίηση RESTful APIs και κατάλληλη τεκμηρίωση',
    'Σχεδιασμός σχεσιακών βάσεων δεδομένων με κατάλληλα ευρετήρια',
    'Εφαρμογή στρατηγικών caching σε επίπεδο εφαρμογής',
    'Εμπειρία πλήρους κύκλου ζωής έργου: από προδιαγραφές έως deployment',
    'Έλεγχος έκδοσης και πρακτικές συνεργατικής ανάπτυξης',
    'Τεκμηρίωση API και πρότυπα Swagger/OpenAPI',
    'Ενοποίηση frontend-backend και διαχείριση ροής δεδομένων'
]
for obj in learning_objectives:
    doc.add_paragraph(obj, style='List Bullet')

# SECTION 3: Project Scope & Constraints
doc.add_heading('3. Scope & Περιορισμοί Έργου', level=1)

doc.add_heading('3.1 Ορισμός Scope', level=2)
scope_points = [
    'Σύστημα μονού χρήστη (προσωπική διαχείριση εργασιών)',
    'Τρεις κύριες οντότητες: Users, Boards, Tasks',
    'Περίπου 10-12 τελικά σημεία API',
    'Μηχανισμός in-memory caching',
    'Ελαφρά βάση δεδομένων (SQLite)',
    'Απλή αλλά λειτουργική διεπαφή χρήστη',
    'Βασική ταυτοποίηση με JWT tokens'
]
for point in scope_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('3.2 Σκόπιμοι Περιορισμοί', level=2)
constraints = [
    'Χωρίς χαρακτηριστικά πραγματικού χρόνου (WebSockets)',
    'Χωρίς κατανεμημένα συστήματα caching',
    'Χωρίς πολύπλοκα μοντέλα δικαιωμάτων',
    'Απομόνωση δεδομένων μονού χρήστη',
    'Χωρίς λειτουργικότητα μεταφόρτωσης αρχείων',
    'Περιορισμένο στη βασική λειτουργικότητα Kanban'
]
for constraint in constraints:
    doc.add_paragraph(constraint, style='List Bullet')

doc.add_heading('3.3 Αιτιολόγηση Σχεδιασμού', level=2)
doc.add_paragraph(
    'Το scope είναι σκόπιμα περιορισμένο για να επιτρέψει τη λεπτομερή κάλυψη βασικών '
    'εννοιών (αρχιτεκτονική, σχεδιασμός βάσης δεδομένων, API, caching) εντός του χρονοδιαγράμματος '
    '3 εβδομάδων, διατηρώντας το δυναμικό για μελλοντικές επεκτάσεις και κλιμακωτότητα.'
)

# SECTION 4: Deliverables
doc.add_heading('4. Παραδοτέα', level=1)

deliverables = [
    ('Τεχνική Τεκμηρίωση', '8-12 σελίδες που καλύπτουν: αρχιτεκτονική, δομές δεδομένων, '
     'ροή δεδομένων, σχήμα βάσης δεδομένων, σχεδιασμός API, σημειώσεις υλοποίησης'),
    ('Πηγαίος Κώδικας', 'Δημόσιο Git repository με συχνά commits, ακολουθώντας τις καλύτερες πρακτικές'),
    ('Λειτουργικό Demo', 'Εκτελέσιμο τοπικά μέσω docker-compose με σαφείς οδηγίες ρύθμισης'),
    ('Τεκμηρίωση API', 'Τεκμηρίωση Swagger/OpenAPI (αυτόματα δημιουργημένη από FastAPI)'),
    ('Σχήμα Βάσης Δεδομένων', 'Πλήρης ER διάγραμμα και τεκμηρίωση σχήματος'),
    ('README & Οδηγός Ρύθμισης', 'Οδηγίες ξεκινήματος για Docker και χειροκίνητη ρύθμιση')
]

for title_del, desc in deliverables:
    doc.add_paragraph(f'{title_del}: {desc}', style='List Bullet')

# SECTION 5: 3-Week Roadmap
doc.add_heading('5. Roadmap 3 Εβδομάδων', level=1)

roadmap_table = doc.add_table(rows=4, cols=3)
roadmap_table.style = 'Light Grid Accent 1'

# Header
header_cells = roadmap_table.rows[0].cells
header_cells[0].text = 'Εβδομάδα'
header_cells[1].text = 'Ώρες'
header_cells[2].text = 'Παραδοτέα & Εργασίες'

# Week 1
week1_cells = roadmap_table.rows[1].cells
week1_cells[0].text = '1'
week1_cells[1].text = '10-12'
week1_cells[2].text = ('Φάση σχεδιασμού: Αρχιτεκτονική, ER διάγραμμα, OpenAPI spec, UI mockups.\n'
                       'Ρύθμιση: Repository, Docker, σχήμα βάσης δεδομένων, JWT auth endpoints.\n'
                       'Πρώτα API endpoints: /auth/register, /auth/login, /auth/me')

# Week 2
week2_cells = roadmap_table.rows[2].cells
week2_cells[0].text = '2'
week2_cells[1].text = '12-14'
week2_cells[2].text = ('Backend: Πλήρης CRUD για Boards και Tasks, JWT middleware, in-memory cache.\n'
                       'Frontend: Φόρμα σύνδεσης, λίστα boards, προβολή board με drag-drop μεταξύ στηλών.\n'
                       'Δοκιμές: Βασικές δοκιμές λειτουργικότητας')

# Week 3
week3_cells = roadmap_table.rows[3].cells
week3_cells[0].text = '3'
week3_cells[1].text = '10-12'
week3_cells[2].text = ('Τελειοποίηση: Endpoint στατιστικών, επαλήθευση εισόδου, χειρισμός σφαλμάτων.\n'
                       'Δοκιμές: Δοκιμές μονάδας και ολοκλήρωσης.\n'
                       'Τεκμηρίωση: README, στιγμιότυπα, έγγραφα Swagger.\n'
                       'Παρουσίαση: Βίντεο demo ή ζωντανό demo')

# Time Management Tips
doc.add_heading('5.1 Συμβουλές Διαχείρισης Χρόνου', level=2)
tips = [
    'Χρησιμοποιήστε SQLite αντί PostgreSQL εάν η ρύθμιση Docker προκαλεί καθυστερήσεις',
    'Χρησιμοποιήστε υπάρχουσες βιβλιοθήκες drag-drop (π.χ. @dnd-kit/core)',
    'Δώστε προτεραιότητα στη βασική λειτουργικότητα έναντι προηγμένων χαρακτηριστικών',
    'Κάντε commits συχνά στο Git (τουλάχιστον καθημερινά)',
    'Γράψτε τεκμηρίωση API καθώς κατασκευάζετε endpoints'
]
for tip in tips:
    doc.add_paragraph(tip, style='List Bullet')

doc.add_page_break()

# SECTION 6: Executive Summary
doc.add_heading('6. Εκτελεστική Περίληψη', level=1)
doc.add_paragraph(
    'Το TaskBoard Lite είναι μια σύγχρονη διαδικτυακή εφαρμογή σχεδιασμένη για να βοηθήσει '
    'τους χρήστες να οργανώσουν και να παρακολουθήσουν εργασίες χρησιμοποιώντας τη μεθοδολογία Kanban. '
    'Η εφαρμογή συνδυάζει ένα backend FastAPI με ένα responsive frontend HTML/CSS/JavaScript, '
    'παρέχοντας μια διαισθητική διεπαφή για διαχείριση εργασιών σε πολλαπλές σανίδες.'
)

# SECTION 7: Technical Architecture
doc.add_heading('7. Τεχνική Αρχιτεκτονική', level=1)

doc.add_heading('7.1 Στοίβα Τεχνολογίας', level=2)
tech_table = doc.add_table(rows=7, cols=2)
tech_table.style = 'Light Grid Accent 1'
cells = tech_table.rows[0].cells
cells[0].text = 'Στοιχείο'
cells[1].text = 'Τεχνολογία'

rows_data = [
    ('Backend Framework', 'FastAPI (Python 3.11)'),
    ('Βάση Δεδομένων', 'SQLite με SQLAlchemy ORM'),
    ('Ταυτοποίηση', 'JWT Tokens (PyJWT)'),
    ('Frontend', 'HTML5, CSS3, Vanilla JavaScript'),
    ('Caching', 'In-memory cache (προσαρμοσμένη υλοποίηση)'),
    ('Τεκμηρίωση API', 'Swagger UI / OpenAPI 3.0')
]
for i, (component, tech) in enumerate(rows_data, 1):
    cells = tech_table.rows[i].cells
    cells[0].text = component
    cells[1].text = tech

doc.add_heading('7.2 Αρχιτεκτονική Συστήματος', level=2)
doc.add_paragraph('Αρχιτεκτονική τριών επιπέδων:')
layers = [
    'Επίπεδο Παρουσίασης: Εφαρμογή ιστού μιας σελίδας (HTML/CSS/JavaScript)',
    'Επίπεδο Επιχειρησιακής Λογικής: RESTful API κατασκευασμένο με FastAPI',
    'Επίπεδο Δεδομένων: Βάση δεδομένων SQLite με SQLAlchemy ORM'
]
for layer in layers:
    doc.add_paragraph(layer, style='List Bullet')

doc.add_heading('7.3 Σχεδιασμός Επιπέδου API', level=2)
doc.add_paragraph('Το API ακολουθεί τις αρχές REST με:')
api_features = [
    'Ανεξάρτητη ταυτοποίηση χρησιμοποιώντας JWT tokens',
    'Συνεπής χειρισμός σφαλμάτων και κωδικοί κατάστασης HTTP',
    'Επαλήθευση αιτήματος/απόκρισης χρησιμοποιώντας σχήματα Pydantic',
    'CORS middleware για επικοινωνία frontend-backend',
    'Αυτόματη τεκμηρίωση API μέσω Swagger'
]
for feature in api_features:
    doc.add_paragraph(feature, style='List Bullet')

# SECTION 8: Data Model
doc.add_heading('8. Μοντέλο Δεδομένων', level=1)

doc.add_heading('8.1 Διάγραμμα Σχέσης Οντοτήτων', level=2)
doc.add_paragraph('Το μοντέλο δεδομένων αποτελείται από τρεις κύριες οντότητες:')

doc.add_heading('8.2 Πίνακας Users', level=2)
user_fields = [
    'id (Integer, Primary Key)',
    'email (String, Unique, Indexed)',
    'username (String)',
    'password_hash (String)',
    'created_at (DateTime, Default: UTC now)'
]
for field in user_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading('8.3 Πίνακας Boards', level=2)
board_fields = [
    'id (Integer, Primary Key)',
    'user_id (Integer, Foreign Key → users.id)',
    'title (String)',
    'description (Text, Optional)',
    'created_at (DateTime, Default: UTC now)'
]
for field in board_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading('8.4 Πίνακας Tasks', level=2)
task_fields = [
    'id (Integer, Primary Key)',
    'board_id (Integer, Foreign Key → boards.id)',
    'title (String)',
    'description (Text, Optional)',
    'status (String: "todo", "in_progress", "done")',
    'position (Integer, Σειρά εντός σανίδας)',
    'created_at (DateTime, Default: UTC now)',
    'updated_at (DateTime, Ενημερώθηκε σε αλλαγές)'
]
for field in task_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading('8.5 Σχέσεις', level=2)
doc.add_paragraph('User ↔ Board: One-to-Many (Ο χρήστης έχει πολλαπλές σανίδες)', style='List Bullet')
doc.add_paragraph('Board ↔ Task: One-to-Many (Η σανίδα περιέχει πολλαπλές εργασίες)', style='List Bullet')

# SECTION 9: API Endpoints
doc.add_heading('9. Τελικά Σημεία API', level=1)

doc.add_heading('9.1 Τελικά Σημεία Ταυτοποίησης', level=2)
auth_endpoints = [
    'POST /auth/register - Εγγραφή νέου χρήστη',
    'POST /auth/login - Ταυτοποίηση χρήστη και λήψη JWT token',
    'GET /auth/me - Λήψη πληροφοριών τρέχοντος ταυτοποιημένου χρήστη'
]
for endpoint in auth_endpoints:
    doc.add_paragraph(endpoint, style='List Bullet')

doc.add_heading('9.2 Τελικά Σημεία Διαχείρισης Σανίδας', level=2)
board_endpoints = [
    'GET /boards - Λίστα όλων των σανίδων για τον τρέχοντα χρήστη',
    'POST /boards - Δημιουργία νέας σανίδας',
    'GET /boards/{board_id} - Λήψη συγκεκριμένης σανίδας',
    'PUT /boards/{board_id} - Ενημέρωση λεπτομερειών σανίδας',
    'DELETE /boards/{board_id} - Διαγραφή σανίδας και σχετικών εργασιών'
]
for endpoint in board_endpoints:
    doc.add_paragraph(endpoint, style='List Bullet')

doc.add_heading('9.3 Τελικά Σημεία Διαχείρισης Εργασιών', level=2)
task_endpoints = [
    'GET /boards/{board_id}/tasks - Λήψη όλων των εργασιών σε σανίδα',
    'POST /boards/{board_id}/tasks - Δημιουργία νέας εργασίας',
    'GET /tasks/{task_id} - Λήψη συγκεκριμένης εργασίας',
    'PUT /tasks/{task_id} - Ενημέρωση εργασίας (τίτλο, περιγραφή, κατάσταση)',
    'DELETE /tasks/{task_id} - Διαγραφή εργασίας'
]
for endpoint in task_endpoints:
    doc.add_paragraph(endpoint, style='List Bullet')

doc.add_heading('9.4 Τελικά Σημεία Στατιστικών', level=2)
stats_endpoints = [
    'GET /stats/boards - Λήψη στατιστικών σανίδας',
    'GET /stats/tasks - Λήψη στατιστικών εργασιών'
]
for endpoint in stats_endpoints:
    doc.add_paragraph(endpoint, style='List Bullet')

doc.add_heading('9.5 Σύνολο Τελικών Σημείων', level=2)
doc.add_paragraph('13 endpoints συνολικά, που καλύπτουν όλη τη βασική λειτουργικότητα')

# SECTION 10: Implementation Status
doc.add_heading('10. Κατάσταση Υλοποίησης', level=1)

doc.add_heading('10.1 Ολοκληρωμένα Στοιχεία', level=2)
completed = [
    '✓ Backend API με όλα τα CRUD endpoints',
    '✓ Ταυτοποίηση χρήστη με JWT tokens',
    '✓ Σχήμα βάσης δεδομένων και μοντέλα SQLAlchemy',
    '✓ Επαλήθευση αιτήματος/απόκρισης (Pydantic)',
    '✓ Μηχανισμός in-memory caching',
    '✓ Φόρμες σύνδεσης και εγγραφής frontend',
    '✓ Διεπαφή διαχείρισης σανίδας και εργασιών',
    '✓ Drag-and-drop κίνηση εργασιών μεταξύ στηλών',
    '✓ Χειρισμός σφαλμάτων και επαλήθευση',
    '✓ Τεκμηρίωση API (Swagger UI)',
    '✓ Διαμόρφωση Docker (Dockerfile, docker-compose.yml)',
    '✓ Οδηγός Ξεκινήματος'
]
for item in completed:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10.2 Παραδοτέα Τεκμηρίωσης', level=2)
docs = [
    'Τεχνική Προδιαγραφή (8-12 σελίδες)',
    'Επιχειρηματική Προδιαγραφή (ολοκληρωμένο έγγραφο)',
    'Οδηγός Ξεκινήματος (οδηγίες ρύθμισης)',
    'Διάγραμμα Δομής Repository (PNG)',
    'Τεκμηρίωση Swagger/OpenAPI (αυτόματα δημιουργημένη)',
    'README.md με επισκόπηση έργου'
]
for doc_item in docs:
    doc.add_paragraph(doc_item, style='List Bullet')

doc.add_heading('10.3 Διαχείριση Repository', level=2)
repo_points = [
    'Δημόσιο GitHub repository: github.com/tsiartas48/taskboard-lite',
    'Συχνά commits με περιγραφικά μηνύματα',
    'Σαφής ιστορία commits που παρακολουθεί την ανάπτυξη',
    'Όλος ο κώδικας και τεκμηρίωση στον έλεγχο έκδοσης'
]
for point in repo_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_page_break()

# SECTION 11: Security Considerations
doc.add_heading('11. Θέσεις Ασφάλειας', level=1)

security = [
    'Ταυτοποίηση JWT: Όλα τα προστατευμένα endpoints απαιτούν έγκυρα tokens',
    'Κρυπτογράφηση Κωδικών: Bcrypt για ασφαλή αποθήκευση κωδικών',
    'Απομόνωση Δεδομένων: Οι χρήστες μπορούν να έχουν πρόσβαση μόνο στα δικά τους δεδομένα',
    'Προστασία CORS: API διαμορφώνεται για συγκεκριμένες προελεύσεις',
    'Επαλήθευση Εισόδου: Τα σχήματα Pydantic επαληθεύουν όλα τα αιτήματα',
    'Χειρισμός Σφαλμάτων: Γενικά μηνύματα σφαλμάτων αποτρέπουν διαρροές πληροφοριών'
]
for sec in security:
    doc.add_paragraph(sec, style='List Bullet')

# SECTION 12: Caching Strategy
doc.add_heading('12. Στρατηγική Caching', level=1)

doc.add_heading('12.1 In-Memory Cache', level=2)
doc.add_paragraph(
    'Η εφαρμογή υλοποιεί ένα απλό in-memory caching layer για να μειώσει τα ερωτήματα βάσης δεδομένων '
    'και να βελτιώσει την απόδοση. Τα κλειδιά cache δημιουργούνται με βάση το user_id για να εξασφαλίσουν '
    'απομόνωση δεδομένων.'
)

doc.add_heading('12.2 Άκυρωση Cache', level=2)
cache_invalidation = [
    'Η cache ακυρώνεται αυτόματα όταν αλλάζουν τα δεδομένα',
    'Η cache σανίδων ακυρώνεται κατά τη δημιουργία/ενημέρωση/διαγραφή σανίδας',
    'Η cache εργασιών ακυρώνεται κατά την αλλαγή εργασιών',
    'Η χειροκίνητη εκκαθάριση cache είναι διαθέσιμη για debugging'
]
for point in cache_invalidation:
    doc.add_paragraph(point, style='List Bullet')

# SECTION 13: Deployment Options
doc.add_heading('13. Επιλογές Ανάπτυξης', level=1)

doc.add_heading('13.1 Ανάπτυξη Docker', level=2)
doc.add_paragraph('Πλήρως δοχειοποιημένο χρησιμοποιώντας docker-compose:')
docker_items = [
    'Υπηρεσία Backend (Python/FastAPI)',
    'Υπηρεσία Frontend (HTML/CSS/JavaScript)',
    'Volume mounts για ανάπτυξη',
    'Health checks για παρακολούθηση υπηρεσίας',
    'Απομόνωση δικτύου χρησιμοποιώντας προσαρμοσμένο δίκτυο'
]
for item in docker_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('13.2 Χειροκίνητη Ανάπτυξη', level=2)
doc.add_paragraph('Μπορεί επίσης να τρέξει τοπικά χωρίς Docker:')
manual_items = [
    'Ρύθμιση περιβάλλοντος Python virtual',
    'Άμεση εκτέλεση uvicorn server',
    'HTTP server για στατικά αρχεία frontend',
    'Τοπική βάση δεδομένων SQLite'
]
for item in manual_items:
    doc.add_paragraph(item, style='List Bullet')

# SECTION 14: Evaluation Criteria
doc.add_heading('14. Κριτήρια Αξιολόγησης', level=1)

eval_table = doc.add_table(rows=3, cols=2)
eval_table.style = 'Light Grid Accent 1'

header_cells = eval_table.rows[0].cells
header_cells[0].text = 'Κριτήριο'
header_cells[1].text = 'Απαιτήσεις'

eval1_cells = eval_table.rows[1].cells
eval1_cells[0].text = 'Έκθεση & Αρχιτεκτονική'
eval1_cells[1].text = 'Πλήρης τεχνική τεκμηρίωση με διαγράμματα αρχιτεκτονικής, αιτιολόγηση σχεδιασμού και σημειώσεις υλοποίησης'

eval2_cells = eval_table.rows[2].cells
eval2_cells[0].text = 'Λειτουργικότητα MVP'
eval2_cells[1].text = 'Όλα τα endpoints λειτουργούν σωστά, το demo εφαρμογής τρέχει, ο χρήστης μπορεί να δημιουργήσει σανίδες/εργασίες και να τις μετακινήσει μεταξύ στηλών'

# SECTION 15: Alternative Domains
doc.add_heading('15. Εναλλακτικά Project Domains', level=1)
doc.add_paragraph(
    'Διατηρώντας την ίδια αρχιτεκτονική (3 πίνακες, ~10-12 endpoints, in-memory cache), '
    'το έργο μπορεί να υλοποιηθεί με διαφορετικά domains:'
)

alternative_domains = [
    'Προσωπική Βιβλιοθήκη: Παρακολούθηση προόδου βιβλίων με αξιολογήσεις και στατιστικά',
    'Tracker Εξόδων: Έξοδα βάσει κατηγορίας με μηνιαίες περιλήψεις και στατιστικά',
    'Habit Tracker: Καθημερινές συνήθειες με παρακολούθηση streak και ποσοστά ολοκλήρωσης',
    'Λίστα Εστιατορίων: Κριτικές εστιατορίων με γεωθέση και σύστημα αξιολόγησης',
    'Χαρτοφυλάκιο Έργων: Προσωπικά έργα με παρακολούθηση προόδου και δεξιότητες'
]
for domain in alternative_domains:
    doc.add_paragraph(domain, style='List Bullet')

# SECTION 16: Future Enhancement Opportunities
doc.add_heading('16. Ευκαιρίες Μελλοντικής Βελτίωσης', level=1)

enhancements = [
    'Συνεργατικές σανίδες: Κοινοποίηση σανίδων με άλλους χρήστες',
    'Προηγμένη φιλτραρίσματα: Φιλτραρίσματα και αναζήτηση εργασιών κατά διάφορα κριτήρια',
    'Ημερομηνίες λήξης και υπενθυμίσεις: Διαχείριση εργασιών βάσει χρόνου',
    'Συνημμένα αρχεία: Προσάρτηση εγγράφων σε εργασίες',
    'Ιστορικό δραστηριότητας: Παρακολούθηση όλων των αλλαγών σανίδων και εργασιών',
    'Εφαρμογή Mobile: Εγγενείς εφαρμογές iOS/Android',
    'Dark Mode: Επιλογές προσαρμογής θέματος',
    'Εξάρτηση Εργασιών: Ορισμός σχέσεων εργασιών και ακολουθίας',
    'Επαναλαμβανόμενες Εργασίες: Αυτοματοποίηση δημιουργίας επαναλαμβανόμενων εργασιών',
    'Ενσωματώσεις: Σύνδεση με εξωτερικές υπηρεσίες (Slack, email, calendar)',
    'Προηγμένα Αναλυτικά: Πίνακας με λεπτομερή μετρικά',
    'Συνεργασία Ομάδας: Πολυ-χρήστη workspaces'
]
for enhancement in enhancements:
    doc.add_paragraph(enhancement, style='List Bullet')

# SECTION 17: Success Metrics
doc.add_heading('17. Μετρικές Επιτυχίας', level=1)

metrics = [
    'Όλα τα API endpoints δοκιμάστηκαν και λειτουργούν σωστά',
    'Το Frontend επικοινωνεί με το Backend με επιτυχία',
    'Ο χρήστης μπορεί να ολοκληρώσει πλήρη ροή εργασίας: εγγραφή, σύνδεση, δημιουργία σανίδας, προσθήκη εργασιών, μετακίνηση εργασιών',
    'Τα δεδομένα αποθηκεύονται σωστά στη βάση δεδομένων',
    'Ο χειρισμός σφαλμάτων εμφανίζει ευανάγνωστα μηνύματα χρήστη',
    'Ο κώδικας είναι καλά οργανωμένος και τεκμηριωμένος',
    'Η ιστορία Git δείχνει συνεπή ανάπτυξη',
    'Η ρύθμιση Docker τρέχει χωρίς σφάλματα',
    'Η τεκμηρίωση API είναι αυτόματα δημιουργημένη και ακριβής'
]
for metric in metrics:
    doc.add_paragraph(metric, style='List Bullet')

# SECTION 18: Conclusion
doc.add_heading('18. Συμπέρασμα', level=1)
doc.add_paragraph(
    'Το TaskBoard Lite αντιπροσωπεύει ένα ολοκληρωμένο εκπαιδευτικό έργο που καλύπτει τον πλήρη κύκλο '
    'ζωής ανάπτυξης λογισμικού. Υλοποιώντας αυτή την εφαρμογή, οι σπουδαστές αποκτούν πρακτική εμπειρία '
    'σε αρχιτεκτονική συστήματος, σχεδιασμό API, μοντελοποίηση βάσης δεδομένων, ανάπτυξη frontend και '
    'σύγχρονες πρακτικές ανάπτυξης.'
)

doc.add_paragraph(
    'Οι σκόπιμοι περιορισμοί και το περιορισμένο scope επιτρέπουν την εις βάθος εξερεύνηση βασικών '
    'εννοιών, ενώ το επεκτάσιμο σχεδιασμό αποδεικνύει ότι υπάρχει δυναμικό για μελλοντική βελτίωση και '
    'σκαλοσιτύ. Ο συνδυασμός τεχνικής τεκμηρίωσης, λειτουργικού κώδικα και παραδειγμάτων ανάπτυξης '
    'παρέχει ένα ολοκληρωμένο portfolio piece που δείχνει δεξιότητα στην ανάπτυξη full-stack.'
)

# Footer
doc.add_paragraph()
footer = doc.add_paragraph('---')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

footer_text = doc.add_paragraph(f'Έγγραφο δημιουργήθηκε στις {datetime.now().strftime("%d %B %Y, %H:%M")}')
footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_text.runs[0]
footer_run.font.size = Pt(9)
footer_run.font.italic = True
footer_run.font.color.rgb = RGBColor(128, 128, 128)

footer_text2 = doc.add_paragraph('Έκδοση 2.0 - Ολοκληρωμένη Προδιαγραφή Έργου & Τεχνική')
footer_text2.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run2 = footer_text2.runs[0]
footer_run2.font.size = Pt(9)
footer_run2.font.italic = True
footer_run2.font.color.rgb = RGBColor(128, 128, 128)

# Save document
output_path = 'TaskBoard_Lite_Tekniki_Ekthesi_Greek.docx'
doc.save(output_path)
print(f'✓ Ελληνικό έγγραφο προδιαγραφών δημιουργήθηκε: {output_path}')
