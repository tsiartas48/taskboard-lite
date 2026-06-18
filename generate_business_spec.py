from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create Document
doc = Document()

# Set up styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title_run = title.add_run('TaskBoard Lite')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Business Specification Document')
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(68, 114, 196)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Date
date_para = doc.add_paragraph(f'Document Date: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'TaskBoard Lite is a personal task management application that combines the efficiency of '
    'Kanban board methodology with modern web technologies. The application enables individual users '
    'to organize, track, and manage their tasks across multiple boards with an intuitive drag-and-drop interface.'
)

# Business Objectives
doc.add_heading('2. Business Objectives', level=1)
objectives = [
    'Provide an efficient task management solution for individual users',
    'Implement Kanban board visualization for better workflow management',
    'Enable secure user authentication and data isolation',
    'Deliver a responsive and user-friendly interface',
    'Ensure data persistence and reliability',
    'Support future scalability and feature expansion'
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

# Target Users
doc.add_heading('3. Target Users', level=1)
doc.add_paragraph('Primary Target Audience:')
users = [
    'Individual professionals and knowledge workers',
    'Students managing academic projects and assignments',
    'Freelancers organizing multiple client tasks',
    'Small team members requiring personal task tracking',
    'Anyone seeking a lightweight task management solution'
]
for user in users:
    doc.add_paragraph(user, style='List Bullet')

# Key Features
doc.add_heading('4. Key Features', level=1)

doc.add_heading('4.1 User Authentication', level=2)
doc.add_paragraph(
    'Secure user registration and login functionality with JWT-based authentication tokens. '
    'Each user\'s data is isolated and protected with industry-standard password hashing.'
)

doc.add_heading('4.2 Board Management', level=2)
doc.add_paragraph(
    'Users can create multiple Kanban boards to organize tasks by project, context, or any custom grouping. '
    'Each board includes a title and optional description for context.'
)

doc.add_heading('4.3 Task Management', level=2)
doc.add_paragraph(
    'Full CRUD operations for tasks including:'
)
features = [
    'Create tasks with title and optional description',
    'Move tasks between status columns (To Do, In Progress, Done, etc.)',
    'Update task details and status',
    'Delete completed or unnecessary tasks',
    'Organize tasks with position-based ordering'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('4.4 Performance Optimization', level=2)
doc.add_paragraph(
    'Built-in caching mechanism to reduce database queries and improve application responsiveness. '
    'Cache is intelligently invalidated when data changes.'
)

doc.add_heading('4.5 Statistics & Analytics', level=2)
doc.add_paragraph(
    'Dashboard providing insights into task completion rates, board activity, and workflow metrics.'
)

# Business Benefits
doc.add_heading('5. Business Benefits', level=1)
benefits = [
    'Improved Productivity: Kanban methodology helps visualize workflow and identify bottlenecks',
    'Better Organization: Centralized location for all personal and professional tasks',
    'Data Privacy: User data is stored securely on private infrastructure',
    'Ease of Use: Intuitive interface requires minimal learning curve',
    'Cost Effective: Lightweight solution with minimal infrastructure requirements',
    'Customizable: Multiple boards allow organization by any preferred taxonomy'
]
for benefit in benefits:
    doc.add_paragraph(benefit, style='List Bullet')

# Technical Architecture Overview
doc.add_heading('6. Technical Architecture', level=1)

doc.add_heading('6.1 Technology Stack', level=2)
tech_table = doc.add_table(rows=6, cols=2)
tech_table.style = 'Light Grid Accent 1'
cells = tech_table.rows[0].cells
cells[0].text = 'Component'
cells[1].text = 'Technology'

rows_data = [
    ('Backend Framework', 'FastAPI (Python)'),
    ('Database', 'SQLite'),
    ('ORM', 'SQLAlchemy'),
    ('Frontend', 'HTML5, CSS3, Vanilla JavaScript'),
    ('Authentication', 'JWT Tokens')
]
for i, (component, tech) in enumerate(rows_data, 1):
    cells = tech_table.rows[i].cells
    cells[0].text = component
    cells[1].text = tech

doc.add_heading('6.2 System Architecture', level=2)
doc.add_paragraph('The system follows a three-tier architecture:')
doc.add_paragraph('Frontend Layer: Single-page web application serving HTML, CSS, and JavaScript', style='List Bullet')
doc.add_paragraph('API Layer: RESTful API built with FastAPI providing all business logic', style='List Bullet')
doc.add_paragraph('Data Layer: SQLite database with SQLAlchemy ORM for data persistence', style='List Bullet')

# Database Schema
doc.add_heading('7. Data Model', level=1)

doc.add_heading('7.1 Users Table', level=2)
doc.add_paragraph('Stores user account information and credentials')
fields = [
    'id (Primary Key)',
    'email (Unique identifier)',
    'username (Display name)',
    'password_hash (Encrypted password)',
    'created_at (Account creation timestamp)'
]
for field in fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading('7.2 Boards Table', level=2)
doc.add_paragraph('Stores Kanban boards created by users')
fields = [
    'id (Primary Key)',
    'user_id (Foreign Key to Users)',
    'title (Board name)',
    'description (Optional board description)',
    'created_at (Board creation timestamp)'
]
for field in fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading('7.3 Tasks Table', level=2)
doc.add_paragraph('Stores individual tasks within boards')
fields = [
    'id (Primary Key)',
    'board_id (Foreign Key to Boards)',
    'title (Task name)',
    'description (Optional task description)',
    'status (Task status: todo, in_progress, done)',
    'position (Order within board)',
    'created_at (Task creation timestamp)',
    'updated_at (Last modification timestamp)'
]
for field in fields:
    doc.add_paragraph(field, style='List Bullet')

# API Endpoints
doc.add_heading('8. API Endpoints Overview', level=1)

doc.add_heading('8.1 Authentication Endpoints', level=2)
auth_endpoints = [
    'POST /auth/register - User registration',
    'POST /auth/login - User login and token generation',
    'GET /auth/me - Get current user information'
]
for endpoint in auth_endpoints:
    doc.add_paragraph(endpoint, style='List Bullet')

doc.add_heading('8.2 Board Endpoints', level=2)
board_endpoints = [
    'GET /boards - Retrieve all boards for current user',
    'POST /boards - Create new board',
    'GET /boards/{board_id} - Get specific board',
    'PUT /boards/{board_id} - Update board',
    'DELETE /boards/{board_id} - Delete board'
]
for endpoint in board_endpoints:
    doc.add_paragraph(endpoint, style='List Bullet')

doc.add_heading('8.3 Task Endpoints', level=2)
task_endpoints = [
    'GET /boards/{board_id}/tasks - Get all tasks in board',
    'POST /boards/{board_id}/tasks - Create new task',
    'GET /tasks/{task_id} - Get specific task',
    'PUT /tasks/{task_id} - Update task details',
    'DELETE /tasks/{task_id} - Delete task'
]
for endpoint in task_endpoints:
    doc.add_paragraph(endpoint, style='List Bullet')

doc.add_heading('8.4 Statistics Endpoints', level=2)
stats_endpoints = [
    'GET /stats/boards - Get board statistics',
    'GET /stats/tasks - Get task statistics'
]
for endpoint in stats_endpoints:
    doc.add_paragraph(endpoint, style='List Bullet')

# Security Considerations
doc.add_heading('9. Security Considerations', level=1)
security = [
    'JWT Authentication: All API endpoints (except registration and login) require valid JWT tokens',
    'Password Security: User passwords are hashed using industry-standard algorithms',
    'Data Isolation: Users can only access their own boards and tasks',
    'CORS Configuration: API accepts requests from frontend domain',
    'Input Validation: All user inputs are validated before processing'
]
for sec in security:
    doc.add_paragraph(sec, style='List Bullet')

# Deployment & Operations
doc.add_heading('10. Deployment & Operations', level=1)

doc.add_heading('10.1 Deployment Architecture', level=2)
doc.add_paragraph(
    'The application can be deployed on any platform supporting Python and modern web browsers. '
    'The lightweight nature of SQLite makes it suitable for single-user or small-team deployments.'
)

doc.add_heading('10.2 Database Backup', level=2)
doc.add_paragraph(
    'Since SQLite stores data in a single file (taskboard.db), backup is as simple as copying the file. '
    'Regular backups are recommended for production deployments.'
)

doc.add_heading('10.3 Performance Monitoring', level=2)
doc.add_paragraph(
    'Application includes built-in caching for frequently accessed data. '
    'Cache statistics can be monitored to optimize performance.'
)

# Future Enhancements
doc.add_heading('11. Future Enhancement Opportunities', level=1)
enhancements = [
    'Collaborative boards: Share boards with team members',
    'Advanced filtering: Filter tasks by various criteria',
    'Due dates and reminders: Add temporal task management',
    'File attachments: Attach files to tasks',
    'Activity history: Track changes to boards and tasks',
    'Mobile app: Native mobile application',
    'Dark mode: Theme customization',
    'Task dependencies: Create task relationships',
    'Recurring tasks: Automate repetitive tasks',
    'Integrations: Connect with other productivity tools'
]
for enhancement in enhancements:
    doc.add_paragraph(enhancement, style='List Bullet')

# Success Metrics
doc.add_heading('12. Success Metrics', level=1)
doc.add_paragraph('The success of TaskBoard Lite will be measured by:')
metrics = [
    'User Adoption Rate: Number of active users and usage frequency',
    'Data Volume: Growth in number of boards and tasks managed',
    'System Performance: Response time and availability metrics',
    'User Satisfaction: Feedback and ratings from users',
    'Error Rate: Reduction in application errors over time'
]
for metric in metrics:
    doc.add_paragraph(metric, style='List Bullet')

# Conclusion
doc.add_heading('13. Conclusion', level=1)
doc.add_paragraph(
    'TaskBoard Lite provides a lightweight, efficient solution for personal task management. '
    'Built on modern, proven technologies, it delivers a balance of functionality, performance, and ease of use. '
    'The application is designed to be intuitive for individual users while maintaining the flexibility '
    'for future enhancement and scalability.'
)

# Footer
doc.add_paragraph()
footer = doc.add_paragraph('---')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

footer_text = doc.add_paragraph(f'Document generated on {datetime.now().strftime("%B %d, %Y at %H:%M")}')
footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_text.runs[0]
footer_run.font.size = Pt(9)
footer_run.font.italic = True
footer_run.font.color.rgb = RGBColor(128, 128, 128)

# Save document
output_path = 'TaskBoard_Lite_Business_Specification.docx'
doc.save(output_path)
print(f'✓ Business specification document created successfully: {output_path}')
